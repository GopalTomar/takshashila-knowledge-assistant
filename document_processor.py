import os
import pytesseract
from PIL import Image
from typing import List, Dict
import PyPDF2
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import logging

class DocumentProcessor:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
    def extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting from PDF {file_path}: {e}")
            return ""
    
    def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            
            return text
        except Exception as e:
            logging.error(f"Error extracting from DOCX {file_path}: {e}")
            return ""
    
    def extract_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            logging.error(f"Error extracting from image {file_path}: {e}")
            return ""
    
    def extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files"""
        try:
            df = pd.read_excel(file_path, sheet_name=None)
            text = ""
            for sheet_name, sheet_df in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_df.to_string(index=False) + "\n\n"
            return text
        except Exception as e:
            logging.error(f"Error extracting from Excel {file_path}: {e}")
            return ""
    
    def extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logging.error(f"Error extracting from TXT {file_path}: {e}")
            return ""
    
    def process_document(self, file_path: str) -> List[Document]:
        """Process a single document and return chunks"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self.extract_from_pdf(file_path)
        elif file_extension == '.docx':
            text = self.extract_from_docx(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            text = self.extract_from_image(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            text = self.extract_from_excel(file_path)
        elif file_extension == '.txt':
            text = self.extract_from_txt(file_path)
        else:
            logging.warning(f"Unsupported file type: {file_extension}")
            return []
        
        if not text.strip():
            logging.warning(f"No text extracted from {file_path}")
            return []
        
        # Split text into chunks
        texts = self.text_splitter.split_text(text)
        
        # Create Document objects
        documents = []
        for i, chunk_text in enumerate(texts):
            doc = Document(
                page_content=chunk_text,
                metadata={
                    "source": file_path,
                    "chunk_id": i,
                    "file_type": file_extension,
                    "total_chunks": len(texts)
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[Document]:
        """Process multiple documents"""
        all_documents = []
        for file_path in file_paths:
            print(f"Processing: {os.path.basename(file_path)}")
            docs = self.process_document(file_path)
            all_documents.extend(docs)
        
        return all_documents